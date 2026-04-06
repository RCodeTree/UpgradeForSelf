from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Helper function to set cell borders
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        if border_name in kwargs:
            edge_el = OxmlElement(f'w:{border_name}')
            edge_el.set(qn('w:val'), kwargs[border_name])
            edge_el.set(qn('w:sz'), '4')  # border size
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), 'auto')
            tcPr.append(edge_el)

def create_detailed_agreement():
    doc = Document()
    
    # 全局样式设置：中文宋体，英文 Times New Roman
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5  # 1.5倍行距，显得不拥挤

    # --- 标题 ---
    title = doc.add_paragraph('学生岗位实习协议书')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(22)
    run.font.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    subtitle = doc.add_paragraph('（双边详尽版）')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    
    doc.add_paragraph('\n')
    p_no = doc.add_paragraph('协议编号：____________________')
    p_no.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # --- 头部信息 ---
    def add_bold_label(p, label, content=""):
        run = p.add_run(label)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p.add_run(content)

    p = doc.add_paragraph()
    add_bold_label(p, '甲方（实习单位）：')
    p.add_run('_________________________________________')
    
    p = doc.add_paragraph()
    add_bold_label(p, '法定代表人/负责人：')
    p.add_run('_______________________ ')
    add_bold_label(p, '联系电话：')
    p.add_run('__________________')

    p = doc.add_paragraph()
    add_bold_label(p, '通讯地址：')
    p.add_run('___________________________________________________')

    doc.add_paragraph('') # 空行

    p = doc.add_paragraph()
    add_bold_label(p, '乙方（实习学生）：')
    p.add_run('_________________________________________')
    
    p = doc.add_paragraph()
    add_bold_label(p, '身份证号：')
    p.add_run('___________________________________________________')
    
    p = doc.add_paragraph()
    add_bold_label(p, '就读学校及专业：')
    p.add_run('_____________________________________________')

    p = doc.add_paragraph()
    add_bold_label(p, '家庭住址：')
    p.add_run('___________________________________________________')

    p = doc.add_paragraph()
    add_bold_label(p, '联系电话：')
    p.add_run('__________________ ')
    add_bold_label(p, '监护人电话：')
    p.add_run('__________________')

    doc.add_paragraph('_' * 45)

    # --- 前言 ---
    intro = doc.add_paragraph('    为规范学生岗位实习工作，提升人才培养质量，维护实习单位和学生的合法权益，根据《中华人民共和国民法典》《职业学校学生实习管理规定》（2021年修订）等相关法律法规。甲、乙双方本着平等自愿、协商一致的原则，就乙方赴甲方进行岗位实习事宜，签订本协议。')
    intro.paragraph_format.first_line_indent = Inches(0.3)
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- 辅助函数：添加带序号的章节 ---
    def add_section_title(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        run = p.runs[0]
        run.font.name = 'Times New Roman'
        run.font.bold = True
        run.font.size = Pt(14)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    def add_item(text, bold_prefix=None):
        p = doc.add_paragraph()
        if bold_prefix:
            run = p.add_run(bold_prefix)
            run.font.bold = True
        p.add_run(text)
        p.paragraph_format.space_after = Pt(6)

    # --- 第一章 ---
    add_section_title('一、实习基本信息')
    add_item(' 1. 实习岗位：__________________________ （由甲方填写）', bold_prefix=None)
    add_item(' 2. 实习地点：__________________________')
    add_item(' 3. 实习时间：______年____月____日 至 ______年____月____日')
    add_item(' 4. 工作时间：甲方实行 ______ 工时制，乙方遵照执行。')
    add_item(' 5. 实习报酬：')
    doc.add_paragraph('    (1) 报酬金额：人民币 _______ 元/月（或 _______ 元/天）。')
    doc.add_paragraph('    (2) 支付方式：□ 银行转账  □ 现金  □ 其他：_______')
    doc.add_paragraph('    (3) 支付时间：每月 ____ 日支付上月报酬。')
    add_item(' 6. 食宿条件：')
    doc.add_paragraph('    (1) 就餐：□ 甲方提供  □ 甲方提供餐补  □ 乙方自理')
    doc.add_paragraph('    (2) 住宿：□ 甲方提供  □ 甲方提供房补  □ 乙方自理')

    # --- 第二章 ---
    add_section_title('二、甲方（实习单位）的权利与义务')
    add_item('：向乙方提供真实有效的单位资质，提供符合法律规定且不损害乙方身心健康的工作环境、生活环境以及安全防护条件。', bold_prefix='1. 基本保障')
    add_item('：严格执行国家及地方安全生产和职业卫生有关规定，制定安全生产事故应急预案，保障乙方实习期间的人身安全和身体健康。对乙方进行安全防护知识、岗位操作规程等教育培训并进行考核。', bold_prefix='2. 安全管理')
    add_item('：甲方承诺在实习期间为乙方投保实习责任保险或商业意外伤害保险。责任保险范围应覆盖实习活动的全过程，投保费用由甲方承担。', bold_prefix='3. 保险责任')
    add_item('：甲方应安排合格的专业人员对乙方实习进行指导，并对乙方在实习期间进行管理。', bold_prefix='4. 带教指导')
    add_item('：依法保障乙方的基本权利，甲方承诺不得有以下情形：', bold_prefix='5. 禁止性条款')
    doc.add_paragraph('    (1) 安排乙方从事《未成年工特殊保护规定》中禁忌从事的劳动；')
    doc.add_paragraph('    (2) 安排乙方到酒吧、夜总会、歌厅、洗浴中心、网吧等营业性娱乐场所实习；')
    doc.add_paragraph('    (3) 通过中介机构或有偿代理组织、安排和管理实习工作；')
    doc.add_paragraph('    (4) 安排乙方从事Ⅲ级强度以上体力劳动或其他有害身心健康的实习；')
    doc.add_paragraph('    (5) 安排乙方从事高空、井下、放射性、有毒、易燃易爆等高风险实习；')
    doc.add_paragraph('    (6) 安排乙方在休息日、法定节假日实习；')
    doc.add_paragraph('    (7) 安排乙方加班和上夜班。')
    add_item('：不得向乙方收取实习押金、培训费、管理费等费用，不得扣押乙方的证件，不得要求乙方提供担保。', bold_prefix='6. 费用规范')
    add_item('：报酬原则上按月及时、足额、直接支付给乙方，支付周期不得超过1个月。', bold_prefix='7. 报酬支付')
    add_item('：实习结束时，甲方应根据乙方表现及实习情况，如实出具《实习鉴定》或实习证明。', bold_prefix='8. 考核鉴定')

    # --- 第三章 ---
    add_section_title('三、乙方（学生）的权利与义务')
    add_item('：乙方在实习期间，须严格遵守国家法律法规，以及甲方的各项规章制度、操作规程和劳动纪律。', bold_prefix='1. 遵纪守法')
    add_item('：严格按照甲方安全规程和操作规范开展工作，爱护甲方设施设备。提高自我保护意识，严禁涉黄、涉赌、涉毒、酗酒及参与其他危险活动。', bold_prefix='2. 安全义务')
    add_item('：认真实习，接受甲方指导人员的指导，完成实习任务。不得擅自离岗、消极怠工、无故拒绝实习。', bold_prefix='3. 服从管理')
    add_item('：保守甲方的商业秘密、技术秘密及其他未公开信息。', bold_prefix='4. 保密义务')
    add_item('：在签订本协议时，乙方应将实习情况告知法定监护人（或家长），并确保监护人知情同意。', bold_prefix='5. 知情同意')
    add_item('：乙方因特殊情况确需中途离开或终止实习的，应提前七日向甲方提出书面申请，并经甲方同意、办妥手续后方可离开。', bold_prefix='6. 变更申请')
    add_item('：乙方认为甲方安排的工作内容违反法律或相关规定的，有权拒绝并投诉。实习期间发生人身伤害，有依法获得赔偿的权利。', bold_prefix='7. 维权权利')

    # --- 第四章 ---
    add_section_title('四、协议解除与违约责任')
    add_item('：经甲、乙双方协商一致，可以解除协议，并以书面形式确认。', bold_prefix='1. 协商解除')
    add_item('：', bold_prefix='2. 单方解除')
    doc.add_paragraph('    (1) 甲方有权解除：乙方严重违反甲方规章制度，或严重失职给甲方造成重大损害的；乙方未经许可擅自离岗连续3日以上的。')
    doc.add_paragraph('    (2) 乙方有权解除：甲方未按约定支付报酬或提供违法的劳动条件的；甲方强迫劳动的；甲方违规安排乙方从事高危、夜班工作的。')
    add_item('：因自然灾害、政策调整等不可抗力致使协议不能履行的，双方可解除协议。', bold_prefix='3. 不可抗力')
    add_item('：任何一方违约均须承担违约责任。甲方未按约定购买保险，致使乙方无法获得保险赔偿的，由甲方依法承担全部赔偿责任。', bold_prefix='4. 违约责任')

    # --- 第五章 ---
    add_section_title('五、附则')
    add_item('1. 本协议一式两份，甲、乙双方各执一份，具有同等法律效力。')
    add_item('2. 本协议自双方签字（盖章）之日起生效，至约定实习期届满或乙方实习结束时终止。')
    add_item('3. 本协议未尽事宜，由甲、乙双方协商解决。')

    doc.add_paragraph('\n' * 2)

    # --- 签字区 ---
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    
    cell_a = table.cell(0, 0)
    p = cell_a.paragraphs[0]
    p.add_run('甲方（实习单位盖章）：').font.bold = True
    cell_a.add_paragraph('\n法定代表人/授权代表签字：\n_______________________')
    cell_a.add_paragraph('日期：______年____月____日')

    cell_b = table.cell(0, 1)
    p = cell_b.paragraphs[0]
    p.add_run('乙方（学生签字）：').font.bold = True
    cell_b.add_paragraph('\n\n乙方监护人（签字）：\n_______________________')
    cell_b.add_paragraph('日期：______年____月____日')

    # Save
    filename = '学生岗位实习协议书_双边详尽版.docx'
    doc.save(filename)
    return filename

# Generate the file
create_detailed_agreement()