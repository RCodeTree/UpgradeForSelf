import re
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 路径配置
base_dir = r"d:\Windows Defalt\桌面\UpgradeForSelf\英语\convert"
source_file = os.path.join(base_dir, "专项.md")
short_answer_md = os.path.join(base_dir, "强哥简答题专项.md")
programming_md = os.path.join(base_dir, "强哥编程题专项.md")
short_answer_docx = os.path.join(base_dir, "强哥简答题专项.docx")
programming_docx = os.path.join(base_dir, "强哥编程题专项.docx")

def extract_content(source_path):
    with open(source_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    short_answer_lines = ["# 强哥简答题专项汇总\n\n"]
    programming_lines = ["# 强哥编程题专项汇总\n\n"]

    current_volume = ""
    current_section = None # 'short' or 'prog' or None

    for line in lines:
        stripped = line.strip()
        
        # 识别卷号
        if stripped.startswith("## **卷"):
            current_volume = stripped.replace("**", "") # 去掉加粗标记，保留文本
            short_answer_lines.append(f"\n{current_volume}\n")
            programming_lines.append(f"\n{current_volume}\n")
            current_section = None
            continue
            
        # 识别分割线
        if stripped.startswith("---"):
            current_section = None
            continue

        # 识别部分标题
        if "名词解释题" in stripped and stripped.startswith("###"):
            current_section = 'short'
            continue # 不把标题写入，或者写入一个统一的子标题
        
        if "应用题" in stripped and stripped.startswith("###"):
            current_section = 'prog'
            continue

        # 根据当前部分添加内容
        if current_section == 'short':
            if stripped: # 只有非空行才处理，或者是保持原样
                short_answer_lines.append(line)
        elif current_section == 'prog':
            if stripped:
                programming_lines.append(line)
                
    return short_answer_lines, programming_lines

def write_md(lines, path):
    with open(path, 'w', encoding='utf-8') as f:
        f.writelines(lines)
    print(f"Generated Markdown: {path}")

def add_markdown_paragraph(doc, text):
    """
    简单的 Markdown 到 Docx 段落转换
    支持：
    - #, ##, ### 标题
    - **Bold** 粗体
    - 代码块 (简单的缩进或 ``` 识别)
    """
    stripped = text.strip()
    
    # 忽略空行，或者添加空段落
    if not stripped:
        return

    # 标题处理
    if stripped.startswith('# '):
        doc.add_heading(stripped[2:], level=1)
        return
    elif stripped.startswith('## '):
        doc.add_heading(stripped[3:], level=2)
        return
    elif stripped.startswith('### '):
        doc.add_heading(stripped[4:], level=3)
        return

    # 代码块标记处理 (```)
    if stripped.startswith('```'):
        return # 忽略标记行，或者作为分界线

    # 普通段落
    p = doc.add_paragraph()
    
    # 简单的粗体解析
    # 将字符串按 ** 分割
    # 例如 "这是 **粗体** 文字" -> ["这是 ", "粗体", " 文字"]
    parts = re.split(r'(\*\*.*?\*\*)', stripped)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)

def create_docx(lines, output_path):
    doc = Document()
    
    # 设置默认字体（可选，需要更多代码支持中文字体设置，这里简化处理）
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    in_code_block = False
    
    for line in lines:
        stripped = line.strip()
        
        # 代码块处理
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            # 代码块内容使用特殊的样式或字体
            p = doc.add_paragraph()
            run = p.add_run(line.rstrip())
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            continue

        add_markdown_paragraph(doc, line)

    doc.save(output_path)
    print(f"Generated Word Doc: {output_path}")

def main():
    if not os.path.exists(source_file):
        print(f"Error: Source file not found: {source_file}")
        return

    short_lines, prog_lines = extract_content(source_file)
    
    # 写入 Markdown
    write_md(short_lines, short_answer_md)
    write_md(prog_lines, programming_md)
    
    # 转换为 Word
    create_docx(short_lines, short_answer_docx)
    create_docx(prog_lines, programming_docx)
    
    print("All tasks completed successfully.")

if __name__ == "__main__":
    main()
