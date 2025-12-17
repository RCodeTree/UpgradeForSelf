import os
import img2pdf
from docx import Document
from docx.shared import Inches, Cm
from PIL import Image

def convert_images(input_dir, output_dir, base_name="output"):
    # Ensure output directory exists
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Get all image files
    image_files = [f for f in os.listdir(input_dir) if f.lower().endswith(('.png', '.jpg', '.jpeg'))]
    # Sort files naturally/alphabetically. Since they are page_001, page_002, they sort correctly by string.
    image_files.sort()
    
    full_paths = [os.path.join(input_dir, f) for f in image_files]
    
    if not full_paths:
        print("No images found in", input_dir)
        return
    
    print(f"Found {len(full_paths)} images.")

    # Convert to PDF
    pdf_path = os.path.join(output_dir, f"{base_name}.pdf")
    print(f"Generating PDF: {pdf_path}")
    try:
        # img2pdf handles images very well without re-encoding if possible
        with open(pdf_path, "wb") as f:
            f.write(img2pdf.convert(full_paths))
        print("PDF generated successfully.")
    except Exception as e:
        print(f"Error generating PDF: {e}")

    # Convert to Word
    docx_path = os.path.join(output_dir, f"{base_name}.docx")
    print(f"Generating Word: {docx_path}")
    try:
        doc = Document()
        # Set margins to be smaller (e.g., 1cm) to maximize image space
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)
            
        # A4 width is 21cm. With 1cm margins left/right, usable width is 19cm.
        # 19cm is approx 7.48 inches.
        
        for img_path in full_paths:
            try:
                # Add picture
                doc.add_picture(img_path, width=Cm(19))
                # Add paragraph/page break if needed, but usually images just flow.
                # Adding a page break after each image ensures one image per page
                # which is usually desired for scanned pages.
                # However, the last image shouldn't strictly need a break, but it doesn't hurt.
                # Actually, doc.add_picture inserts it into the current paragraph or creates a new one.
                # Let's assume one image per page is best for readability of slides/pages.
                # doc.add_page_break()
            except Exception as e:
                print(f"Error adding image {img_path} to docx: {e}")
        
        doc.save(docx_path)
        print("Word document generated successfully.")
    except Exception as e:
        print(f"Error generating Word doc: {e}")

if __name__ == "__main__":
    input_dir = r"d:\Windows Defalt\桌面\UpgradeForSelf\Script\spider\chaoxing_images"
    output_dir = r"d:\Windows Defalt\桌面\UpgradeForSelf\Script\convert"
    convert_images(input_dir, output_dir, base_name="chaoxing_converted")
